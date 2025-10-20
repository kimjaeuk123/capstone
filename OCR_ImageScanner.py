import easyocr
import sys
import os
import re
import pandas as pd
from PIL import Image
from PIL import ImageOps, ImageFilter, ImageEnhance
import numpy as np
# import cv2
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import getAscent, getDescent


# EasyOCR 리더 설정
reader = easyocr.Reader(['ko', 'en'])

# 사용자에게 디렉토리 경로를 선택하도록 하는 함수
def select_directory():
    folder_path = filedialog.askdirectory(title="스캔 이미지 폴더 선택")
    return folder_path

output_file_excel = ""
output_file_docx = ""
output_file_pdf = ""

# 결과를 저장할 빈 리스트 생성
data = []

# 글꼴 파일 경로
def get_font_path():
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, 'NanumGothic.ttf')

font_path = get_font_path()


def resize_image(image, scale_factor=3.0):
    width, height = image.size
    new_size = (int(width * scale_factor), int(height * scale_factor))
    return image.resize(new_size, Image.LANCZOS)


def preprocess_image(image):
    return image.point(lambda p: p > 128 and 255)  # 이진화


def enhance_image(image):
    return image.point(lambda p: p * 1.5)  # 대비 향상

def auto_gamma(img, clamp=(0.5, 1.5)):
    # img: L 모드(그레이) 기준
    arr = np.asarray(img, dtype=np.float32)
    mean = arr.mean() / 255.0 + 1e-6  # 0으로 나눔 방지
    target = 0.5                      # 중간톤 목표
    gamma = np.log(target) / np.log(mean)
    # 과한 보정을 피하도록 감마 범위 클램프
    gamma = max(clamp[0], min(clamp[1], gamma))

    lut = [int(((i/255.0) ** gamma) * 255 + 0.5) for i in range(256)]
    return img.point(lut)


def enhance_image_for_ocr(image):
    # 그레이스케일
    g = image.convert('L')

    #배경 밝기 추정 후, 어두우면 반전
    arr = np.array(g)
    corners = np.hstack([
        arr[:50, :50].ravel(),
        arr[:50, -50:].ravel(),
        arr[-50:, :50].ravel(),
        arr[-50:, -50:].ravel()
    ])
    if np.median(corners) < 110:
        g = ImageOps.invert(g)

    #동적 감마 보정
    g = auto_gamma(g, clamp=(0.6, 1.4))

    # autocontrast
    g = ImageOps.autocontrast(g, cutoff=2)

    # 살짝 선명
    g = g.filter(ImageFilter.UnsharpMask(radius=1.2, percent=180, threshold=2))

    return g


# 텍스트 블록들을 좌표 기반으로 자동 줄바꿈 처리(텍스트 그룹화)
def group_text_by_lines(text_data):
    lines = []
    current_line = []
    current_y = None
    threshold = 10

    for result in text_data:
        bbox, text, confidence = result[0], result[1], result[2]
        y = bbox[0][1]

        if current_y is None:
            current_y = y
            current_line.append(text)
        else:
            if abs(y - current_y) < threshold:
                current_line.append(text)
            else:
                lines.append(" ".join(current_line))
                current_line = [text]
                current_y = y

    if current_line:
        lines.append(" ".join(current_line))

    return lines


# 텍스트 추출
def extract_text_from_images(path_dir):
    global data  # data 리스트를 전역 변수로 사용
    file_list = os.listdir(path_dir)

    for file_name in file_list:
        if file_name.lower() in ["output.xlsx", "output.docx"]:
            continue

        file_path = os.path.join(path_dir, file_name)

        try:
            img = Image.open(file_path)
            img = resize_image(img)
            img = enhance_image_for_ocr(img)
            img_cv = np.array(img)

            # 왜곡 보정 제거 → 바로 OCR 실행
            text_data = reader.readtext(img_cv, detail=1)
            lines = group_text_by_lines(text_data)

            data.append({'내용': "\n".join(lines), '이미지': file_path})

        except IOError:
            print(f"파일을 열 수 없습니다: {file_path}")


# 엑셀 및 DOCX 파일로 저장
def save_output_files(path_dir, file_type="all"):
    global data  # data 리스트를 전역 변수로 사용
    global output_file_excel, output_file_docx, output_file_pdf  # 전역 변수로 파일 경로 사용

    # 엑셀 파일 저장
    if file_type in ["all", "excel"]:
        df = pd.DataFrame(data)

        output_file_excel = os.path.join(path_dir, 'output.xlsx')

        # 각 텍스트 줄을 별도의 열로 저장 (공백을 기준으로 나누기)
        split_data = []

        for entry in data:
            content = entry['내용'].splitlines()  # 텍스트를 줄 단위로 분리
            split_row = [line.split() for line in content]  # 각 줄 공백 기준 split
            split_data.extend(split_row)

        if split_data:
            max_columns = max(len(row) for row in split_data)
            for row in split_data:
                row.extend([''] * (max_columns - len(row)))
            df_split = pd.DataFrame(split_data)
        else:
            df_split = pd.DataFrame()

        df_split.to_excel(output_file_excel, index=False, header=False)

    # DOCX 파일 저장
    if file_type in ["all", "docx"]:
        df = pd.DataFrame(data)
        output_file_docx = os.path.join(path_dir, 'output.docx')
        doc = Document()
        doc.add_heading('제목을 적어주세요.', 0)

        for index, row in df.iterrows():
            doc.add_paragraph(row['내용'])

        doc.save(output_file_docx)

    # PDF 파일 저장 (경로만 세팅, 실제 저장은 미리보기/편집에서)
    if file_type in ["all", "pdf"]:
        df = pd.DataFrame(data)
        output_file_pdf = os.path.join(path_dir, 'output.pdf')

    print(f"{file_type} 파일 저장 완료")
    return df  # df 반환


def add_text_to_existing_file(file_type, df):
    if file_type == "엑셀":
        file_path = filedialog.askopenfilename(
            title="엑셀 파일 선택",
            filetypes=[("엑셀 파일", "*.xlsx;*.xls")]
        )
    elif file_type == "DOCX":
        file_path = filedialog.askopenfilename(
            title="DOCX 파일 선택",
            filetypes=[("Word 파일", "*.docx")]
        )

    if not file_path:
        return

    if file_type == "엑셀":
        try:
            existing_df = pd.read_excel(file_path, header=None, engine='openpyxl')
            new_data = df[['내용']]
            split_data = []
            for entry in new_data['내용']:
                content = entry.splitlines()
                for line in content:
                    split_row = line.split()
                    split_data.append(split_row)

            if split_data:
                max_columns = max(len(row) for row in split_data)
                for row in split_data:
                    row.extend([''] * (max_columns - len(row)))
                combined_df = pd.concat([existing_df, pd.DataFrame(split_data)], ignore_index=True)
            else:
                combined_df = existing_df

            combined_df.to_excel(file_path, index=False, header=False)
            messagebox.showinfo("성공", "엑셀 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"엑셀 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)

    elif file_type == "DOCX":
        try:
            existing_doc = Document(file_path)
            for index, row in df.iterrows():
                existing_doc.add_paragraph(row['내용'])
            existing_doc.save(file_path)
            messagebox.showinfo("성공", "DOCX 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"DOCX 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)


def open_file(file_path):
    try:
        os.startfile(file_path)
    except Exception as e:
        messagebox.showerror("파일 열기 오류", f"파일을 열 수 없습니다: {e}")

font_size_map = []  # 선택 영역에 따른 폰트 크기 설정 정보

def save_to_pdf_with_custom_fonts(text_widget, font_map, pdf_path):
    c = canvas.Canvas(pdf_path, pagesize=letter)
    pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
    width, height = letter
    margin = 50
    y_position = height - margin

    index = "1.0"
    while True:
        if text_widget.compare(index, ">=", "end"):
            break

        line_text = text_widget.get(index, f"{index} +1line")

        match = re.match(r"\[이미지:\s*(.*?),\s*(\d+),\s*(\d+)\]", line_text.strip())
        if match:
            image_path, width_str, height_str = match.groups()
            img_width, img_height = float(width_str), float(height_str)

            if y_position - img_height < margin:
                c.showPage()
                y_position = height - margin

            c.drawImage(image_path.strip(), margin, y_position - img_height, width=img_width, height=img_height)
            y_position -= img_height + 10
            index = text_widget.index(f"{index} +1line")
            continue

        if not line_text.strip():
            y_position -= 20
            index = text_widget.index(f"{index} +1line")
            continue

        char_index = index
        max_ascent = 0
        max_descent = 0

        while True:
            char = text_widget.get(char_index)
            if char == "\n":
                break
            tags = text_widget.tag_names(char_index)
            font_size = 12
            for tag in tags:
                if tag.startswith("font_"):
                    try:
                        font_size = int(tag.split("_")[-1])
                    except:
                        pass
            ascent = pdfmetrics.getAscent('NanumGothic') / 1000 * font_size
            descent = pdfmetrics.getDescent('NanumGothic') / 1000 * font_size
            max_ascent = max(max_ascent, ascent)
            max_descent = max(max_descent, descent)
            char_index = text_widget.index(f"{char_index} +1c")

        line_spacing = max_ascent + max_descent + 5

        if y_position - line_spacing < margin:
            c.showPage()
            y_position = height - margin

        baseline = y_position - max_ascent

        x_cursor = margin
        char_index = index
        while True:
            char = text_widget.get(char_index)
            if char == "\n":
                break

            tags = text_widget.tag_names(char_index)
            font_size = 12
            for tag in tags:
                if tag.startswith("font_"):
                    try:
                        font_size = int(tag.split("_")[-1])
                    except:
                        pass

            ascent = pdfmetrics.getAscent('NanumGothic') / 1000 * font_size
            descent = pdfmetrics.getDescent('NanumGothic') / 1000 * font_size
            char_width = pdfmetrics.stringWidth(char, 'NanumGothic', font_size)

            if x_cursor + char_width > width - margin:
                x_cursor = margin
                y_position -= line_spacing
                baseline = y_position - max_ascent
                if y_position - line_spacing < margin:
                    c.showPage()
                    y_position = height - margin
                    baseline = y_position - max_ascent

            c.setFont('NanumGothic', font_size)
            c.drawString(x_cursor, baseline, char)

            x_cursor += char_width
            char_index = text_widget.index(f"{char_index} +1c")

        y_position -= line_spacing
        index = text_widget.index(f"{index} +1line")

    c.save()


def preview_and_edit_pdf():
    root = tk.Tk()
    root.title("PDF 미리보기 및 수정")
    root.geometry("600x650")

    text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=25)
    text_area.pack(pady=10, padx=10)

    extracted_text = "\n\n".join([entry['내용'] for entry in data])
    text_area.insert(tk.END, extracted_text)

    search_frame = tk.Frame(root)
    search_frame.pack(pady=5)

    tk.Label(search_frame, text="검색어:").pack(side=tk.LEFT)
    search_entry = tk.Entry(search_frame, width=20)
    search_entry.pack(side=tk.LEFT, padx=5)

    def apply_font_to_selection():
        try:
            font_size = int(font_size_entry.get())
            start = text_area.index(tk.SEL_FIRST)
            end = text_area.index(tk.SEL_LAST)
            font_size_map.append((start, end, font_size))
            tag_name = f"font_{font_size}"
            text_area.tag_add(tag_name, start, end)
            text_area.tag_config(tag_name, font=("NanumGothic", font_size))
        except tk.TclError:
            messagebox.showerror("오류", "텍스트를 먼저 선택해주세요.")
        except ValueError:
            messagebox.showerror("입력 오류", "숫자 형식의 폰트 크기를 입력해주세요.")

    button_apply_font = tk.Button(root, text="선택 텍스트 크기 적용", command=apply_font_to_selection)
    button_apply_font.pack(pady=5)

    def search_text():
        text_area.tag_remove("highlight", "1.0", tk.END)
        keyword = search_entry.get()
        if not keyword:
            return
        start = "1.0"
        while True:
            start = text_area.search(keyword, start, stopindex=tk.END)
            if not start:
                break
            end = f"{start}+{len(keyword)}c"
            text_area.tag_add("highlight", start, end)
            start = end
        text_area.tag_config("highlight", background="yellow", foreground="black")

    search_button = tk.Button(search_frame, text="검색", command=search_text)
    search_button.pack(side=tk.LEFT)

    font_size_frame = tk.Frame(root)
    font_size_frame.pack()

    tk.Label(font_size_frame, text="폰트 크기:").pack(side=tk.LEFT)
    font_size_entry = tk.Entry(font_size_frame, width=5)
    font_size_entry.insert(0, "12")
    font_size_entry.pack(side=tk.LEFT)

    def save_edited_pdf():
        try:
            font_size = int(font_size_entry.get())
        except ValueError:
            messagebox.showerror("입력 오류", "올바른 숫자 형식의 폰트 크기를 입력해주세요.")
            return

        edited_text = text_area.get("1.0", tk.END)
        new_data = [{'내용': edited_text.strip()}]
        save_to_pdf_with_custom_fonts(text_area, font_size_map, output_file_pdf)
        messagebox.showinfo("저장 완료", "수정된 내용이 PDF 파일에 저장되었습니다.")
        open_file(output_file_pdf)

    def add_image_to_pdf():
        file_path = filedialog.askopenfilename(title="이미지 파일 선택", filetypes=[("이미지 파일", "*.png;*.jpg;*.jpeg")])
        if not file_path:
            return

        size_window = tk.Toplevel(root)
        size_window.title("이미지 크기 조정")

        tk.Label(size_window, text="너비:").pack()
        width_entry = tk.Entry(size_window)
        width_entry.pack()

        tk.Label(size_window, text="높이:").pack()
        height_entry = tk.Entry(size_window)
        height_entry.pack()

        def insert_image():
            width = width_entry.get()
            height = height_entry.get()
            text_area.insert(tk.INSERT, f"\n[이미지: {file_path}, {width}, {height}]\n")
            size_window.destroy()

        tk.Button(size_window, text="추가", command=insert_image).pack()

    button_save = tk.Button(root, text="PDF로 저장", command=save_edited_pdf)
    button_save.pack(pady=10)

    button_add_image = tk.Button(root, text="이미지 추가", command=add_image_to_pdf)
    button_add_image.pack(pady=10)

    root.mainloop()


def show_file_selection_gui():
    root = tk.Tk()
    root.title("원하는 항목 선택")
    root.geometry("300x400")

    def save_excel_file():
        df = save_output_files(path_dir, file_type="excel")
        open_file(output_file_excel)

    def save_docx_file():
        df = save_output_files(path_dir, file_type="docx")
        open_file(output_file_docx)

    def save_pdf_file():
        save_output_files(path_dir, file_type="pdf")
        preview_and_edit_pdf()

    button_excel = tk.Button(root, text="EXCEL 파일 열기 및 저장", command=save_excel_file)
    button_excel.pack(pady=10)

    button_docx = tk.Button(root, text="DOCX 파일 열기 및 저장", command=save_docx_file)
    button_docx.pack(pady=10)

    button_pdf = tk.Button(root, text="PDF 파일 미리보기 및 수정", command=save_pdf_file)
    button_pdf.pack(pady=10)

    button_add_excel = tk.Button(root, text="기존 EXCEL 파일에 텍스트 추가",
                                 command=lambda: add_text_to_existing_file("엑셀", pd.DataFrame(data)))
    button_add_excel.pack(pady=10)

    button_add_docx = tk.Button(root, text="기존 DOCX 파일에 텍스트 추가",
                                command=lambda: add_text_to_existing_file("DOCX", pd.DataFrame(data)))
    button_add_docx.pack(pady=10)

    root.mainloop()


path_dir = select_directory()
if path_dir:
    extract_text_from_images(path_dir)
    show_file_selection_gui()
else:
    messagebox.showerror("경로 오류", "유효한 폴더 경로를 선택해주세요.")

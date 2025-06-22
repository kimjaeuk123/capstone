import easyocr
import sys
import os
import pandas as pd
from PIL import Image
import numpy as np
import cv2
from docx import Document
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# EasyOCR 리더 설정
reader = easyocr.Reader(['ko', 'en'])

# 사용자에게 디렉토리 경로를 선택하도록 하는 함수
def select_directory():
    folder_path = filedialog.askdirectory(title="스캔 이미지 폴더 선택")
    return folder_path

output_file_excel = ""
output_file_docx = ""
output_file_pdf = ""

# 결과를 저장할 빈 리스트 생성
data = []

# 글꼴 파일 경로
def get_font_path():
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.dirname(__file__)
    return os.path.join(base_path, 'NanumGothic.ttf')

font_path = get_font_path()

def save_to_pdf(data, pdf_file_path, font_size=12):  # 기본값 12
    c = canvas.Canvas(pdf_file_path, pagesize=letter)
    pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
    c.setFont('NanumGothic', font_size)

    width, height = letter
    margin = 50
    y_position = height - margin

    for entry in data:
        for line in entry['내용'].splitlines():
            if line.startswith("[이미지:"):
                parts = line[8:-1].split(", ")
                image_path = parts[0]
                img_width = float(parts[1]) if len(parts) > 1 else 0
                img_height = float(parts[2]) if len(parts) > 2 else 0

                if img_width > 0 and img_height > 0:
                    if y_position - img_height < margin:
                        c.showPage()
                        c.setFont('NanumGothic', font_size)
                        y_position = height - margin

                    c.drawImage(image_path, margin, y_position - img_height, width=img_width, height=img_height)
                    y_position -= img_height + 10
            else:
                c.drawString(margin, y_position, line)
                y_position -= font_size + 8  # 폰트 크기에 따라 줄 간격 조절
                if y_position < margin:
                    c.showPage()
                    c.setFont('NanumGothic', font_size)
                    y_position = height - margin

    c.save()


# 텍스트 추출
def extract_text_from_images(path_dir):
    global data
    file_list = os.listdir(path_dir)

    for file_name in file_list:
        if file_name.lower() in ["output.xlsx", "output.docx"]:
            continue

        file_path = os.path.join(path_dir, file_name)

        try:
            img = Image.open(file_path)
            img_cv = np.array(img)

            img_corrected = img_cv

            text_data = reader.readtext(img_corrected, detail=0)  # detail=0: 텍스트만 반환
            joined_text = "\n".join(text_data)  # 줄로 묶기

            data.append({'내용': joined_text, '이미지': file_path})

        except IOError:
            print(f"파일을 열 수 없습니다: {file_path}")



# 엑셀 및 DOCX 파일로 저장
def save_output_files(path_dir, file_type="all"):
    global data
    global output_file_excel, output_file_docx, output_file_pdf

    # 엑셀 파일 저장 - 단어 분할 제거 버전
    if file_type in ["all", "excel"]:
        df = pd.DataFrame(data)
        output_file_excel = os.path.join(path_dir, 'output.xlsx')

        # '내용' 열만 추출해서 한 줄당 한 셀에 저장
        simple_df = df[['내용']]
        simple_df.to_excel(output_file_excel, index=False)  # 헤더 포함 저장

    # DOCX 저장은 그대로 유지
    if file_type in ["all", "docx"]:
        df = pd.DataFrame(data)
        output_file_docx = os.path.join(path_dir, 'output.docx')
        doc = Document()
        doc.add_heading('제목을 적어주세요.', 0)
        for index, row in df.iterrows():
            doc.add_paragraph(row['내용'])
        doc.save(output_file_docx)

    # PDF 저장도 그대로
    if file_type in ["all", "pdf"]:
        df = pd.DataFrame(data)
        output_file_pdf = os.path.join(path_dir, 'output.pdf')
        save_to_pdf(data, output_file_pdf)

    print(f"{file_type} 파일 저장 완료")
    return df



def add_text_to_existing_file(file_type, df):
    # 파일 선택
    if file_type == "엑셀":
        file_path = filedialog.askopenfilename(
            title="엑셀 파일 선택",
            filetypes=[("엑셀 파일", "*.xlsx;*.xls")]
        )
    elif file_type == "DOCX":
        file_path = filedialog.askopenfilename(
            title="DOCX 파일 선택",
            filetypes=[("Word 파일", "*.docx")]
        )

    if not file_path:
        return

    if file_type == "엑셀":
        try:
            existing_df = pd.read_excel(file_path, header=None, engine='openpyxl')

            new_data = df[['내용']]
            split_data = []

            # 각 텍스트 블록을 줄바꿈 포함한 하나의 셀로 저장
            for entry in new_data['내용']:
                content = entry.strip()
                if content:
                    split_data.append([content])  # 한 셀에 저장

            new_df = pd.DataFrame(split_data)
            combined_df = pd.concat([existing_df, new_df], ignore_index=True)

            combined_df.to_excel(file_path, index=False, header=False)
            messagebox.showinfo("성공", "엑셀 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"엑셀 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)

    elif file_type == "DOCX":
        try:
            existing_doc = Document(file_path)
            for index, row in df.iterrows():
                existing_doc.add_paragraph(row['내용'])
            existing_doc.save(file_path)
            messagebox.showinfo("성공", "DOCX 파일에 텍스트가 추가되었습니다.")
            open_file(file_path)
        except Exception as e:
            messagebox.showerror("오류", f"DOCX 파일 처리 중 오류가 발생했습니다: {e}")
            print("오류 발생:", e)




# 파일 열기 함수
def open_file(file_path):
    try:
        os.startfile(file_path)
    except Exception as e:
        messagebox.showerror("파일 열기 오류", f"파일을 열 수 없습니다: {e}")


# PDF 미리보기 및 수정 GUI
def preview_and_edit_pdf():
    root = tk.Tk()
    root.title("PDF 미리보기 및 수정")
    root.geometry("600x650")

    text_area = scrolledtext.ScrolledText(root, wrap=tk.WORD, width=70, height=25)
    text_area.pack(pady=10, padx=10)

    extracted_text = "\n\n".join([entry['내용'] for entry in data])
    text_area.insert(tk.END, extracted_text)

    # 검색 기능 UI 추가
    search_frame = tk.Frame(root)
    search_frame.pack(pady=5)

    tk.Label(search_frame, text="검색어:").pack(side=tk.LEFT)
    search_entry = tk.Entry(search_frame, width=20)
    search_entry.pack(side=tk.LEFT, padx=5)

    def search_text():
        text_area.tag_remove("highlight", "1.0", tk.END)  # 기존 하이라이트 제거
        keyword = search_entry.get()
        if not keyword:
            return
        start = "1.0"
        while True:
            start = text_area.search(keyword, start, stopindex=tk.END)
            if not start:
                break
            end = f"{start}+{len(keyword)}c"
            text_area.tag_add("highlight", start, end)
            start = end
        text_area.tag_config("highlight", background="yellow", foreground="black")

    search_button = tk.Button(search_frame, text="검색", command=search_text)
    search_button.pack(side=tk.LEFT)

    #폰트 크기 입력 필드 추가
    font_size_frame = tk.Frame(root)
    font_size_frame.pack()

    tk.Label(font_size_frame, text="폰트 크기:").pack(side=tk.LEFT)
    font_size_entry = tk.Entry(font_size_frame, width=5)
    font_size_entry.insert(0, "12")  # 기본값 12
    font_size_entry.pack(side=tk.LEFT)

    def save_edited_pdf():
        try:
            font_size = int(font_size_entry.get())
        except ValueError:
            messagebox.showerror("입력 오류", "올바른 숫자 형식의 폰트 크기를 입력해주세요.")
            return

        edited_text = text_area.get("1.0", tk.END)
        new_data = [{'내용': edited_text.strip()}]
        save_to_pdf(new_data, output_file_pdf, font_size=font_size)
        messagebox.showinfo("저장 완료", "수정된 내용이 PDF 파일에 저장되었습니다.")
        open_file(output_file_pdf)

    def add_image_to_pdf():
        file_path = filedialog.askopenfilename(title="이미지 파일 선택", filetypes=[("이미지 파일", "*.png;*.jpg;*.jpeg")])
        if not file_path:
            return

        size_window = tk.Toplevel(root)
        size_window.title("이미지 크기 조정")

        tk.Label(size_window, text="너비:").pack()
        width_entry = tk.Entry(size_window)
        width_entry.pack()

        tk.Label(size_window, text="높이:").pack()
        height_entry = tk.Entry(size_window)
        height_entry.pack()

        def insert_image():
            width = width_entry.get()
            height = height_entry.get()
            text_area.insert(tk.INSERT, f"\n[이미지: {file_path}, {width}, {height}]\n")
            size_window.destroy()

        tk.Button(size_window, text="추가", command=insert_image).pack()

    button_save = tk.Button(root, text="PDF로 저장", command=save_edited_pdf)
    button_save.pack(pady=10)

    button_add_image = tk.Button(root, text="이미지 추가", command=add_image_to_pdf)
    button_add_image.pack(pady=10)

    root.mainloop()



# 파일 선택 GUI
def show_file_selection_gui():
    root = tk.Tk()
    root.title("원하는 항목 선택")
    root.geometry("300x400")

    # 엑셀 파일 열기 버튼
    def save_excel_file():
        df = save_output_files(path_dir, file_type="excel")  # 엑셀 파일만 저장
        open_file(output_file_excel)

    # DOCX 파일 열기 버튼
    def save_docx_file():
        df = save_output_files(path_dir, file_type="docx")  # DOCX 파일만 저장
        open_file(output_file_docx)

    # PDF 파일 미리보기 및 저장 함수
    def save_pdf_file():
        save_output_files(path_dir, file_type="pdf")  # PDF 파일만 저장
        preview_and_edit_pdf()  # PDF 미리보기 및 수정

    # 각 버튼 클릭 시 실행될 함수들
    button_excel = tk.Button(root, text="엑셀 파일 열기", command=save_excel_file)  # 엑셀 파일만 열도록 수정
    button_excel.pack(pady=10)

    button_docx = tk.Button(root, text="DOCX 파일 열기", command=save_docx_file)
    button_docx.pack(pady=10)

    button_pdf = tk.Button(root, text="PDF 파일 미리보기 및 수정", command=save_pdf_file)
    button_pdf.pack(pady=10)

    # 엑셀 파일에 텍스트 추가 버튼
    button_add_excel = tk.Button(root, text="기존 엑셀 파일에 텍스트 추가",
                                 command=lambda: add_text_to_existing_file("엑셀", pd.DataFrame(data)))
    button_add_excel.pack(pady=10)

    # DOCX 파일에 텍스트 추가 버튼
    button_add_docx = tk.Button(root, text="기존 DOCX 파일에 텍스트 추가",
                                command=lambda: add_text_to_existing_file("DOCX", pd.DataFrame(data)))
    button_add_docx.pack(pady=10)

    root.mainloop()


# 경로 선택 후 디렉토리에서 이미지 텍스트 추출
path_dir = select_directory()
if path_dir:
    extract_text_from_images(path_dir)
    # GUI에서 버튼 클릭 시 각 파일을 생성하도록 하므로, save_output_files는 여기서 호출되지 않음
    show_file_selection_gui()
else:
    messagebox.showerror("경로 오류", "유효한 폴더 경로를 선택해주세요.")
